# -*- coding: utf-8 -*-
import os, re, json
from datetime import datetime
from docx import Document
from docx.shared import Pt

# ===== CONFIG =====
MODELO_LP = "LP MODELO.docx"
MODELO_EVS = "EVS MODELO.docx"
ARQ_JSON  = "input.json"
PASTA_OUT = "out"
# ==================

def sanitize_filename(name: str) -> str:
    # troca caracteres inválidos e normaliza espaços
    name = re.sub(r'[\\/*?:"<>|]', "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name

def normalize_data(s: str) -> str:
    s = (s or "").strip()
    for fmt in ("%d.%m.%y", "%d.%m.%Y", "%d/%m/%y", "%d/%m/%Y"):
        try:
            return datetime.strptime(s, fmt).strftime("%d.%m.%y")
        except:
            pass
    return s

def inferir_laudo_final(obs: str) -> str:
    t = (obs or "").upper()
    if "REPROVADO" in t or "REPROVADOS" in t: return "(R) REPROVADO"
    if "APROVADO" in t or "APROVADOS" in t:   return "(A) APROVADO"
    return "(A) APROVADO"

def substituir_texto_str(texto: str, mapa: dict) -> str:
    for k, v in mapa.items():
        texto = texto.replace(f"{{{{ {k} }}}}", v)
    return texto

def aplicar_substituicoes(doc: Document, mapa: dict):
    # Parágrafos (não mexe em fonte/tamanho aqui — preserva títulos/cabeçalhos visuais)
    for p in doc.paragraphs:
        novo = substituir_texto_str(p.text, mapa)
        if novo != p.text:
            p.text = novo

    # Tabelas (aplica padronização de fonte)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                novo = substituir_texto_str(c.text, mapa)
                if novo != c.text:
                    c.text = novo
                for par in c.paragraphs:
                    for run in par.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(9 if mapa.get("DATA") and mapa["DATA"] in run.text else 10)

    # Cabeçalhos/Rodapés (substituição, sem alterar fonte)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            if part:
                for p in part.paragraphs:
                    novo = substituir_texto_str(p.text, mapa)
                    if novo != p.text:
                        p.text = novo

def montar_nome_longo(tipo: str, n: str, equipamento: str, cliente: str, solicitante: str, data: str) -> str:
    # Inclui SOLICITANTE na nomenclatura
    return f"REL {tipo.upper()} {n} 26 {equipamento} {cliente} {solicitante} {data}.docx"

def processar_registro(rec: dict) -> dict:
    tipo = (rec.get("TIPO") or "").strip().upper()
    if tipo not in ("LP", "EVS"):
        raise ValueError(f"TIPO inválido: {tipo}. Use 'LP' ou 'EVS'.")

    # permite espaço para casos "693 694"
    n = (rec.get("N_RELATORIO") or "").strip()
    n = re.sub(r"[^\d ]", "", n).strip()
    if not n:
        raise ValueError("N_RELATORIO vazio/inválido.")

    mapa = {
        "TIPO": tipo,
        "N_RELATORIO": n,  # sem /25 (o modelo já traz /25)
        "OBRA_OS_N": (rec.get("OBRA_OS_N") or "").strip(),
        "EQUIPAMENTO": (rec.get("EQUIPAMENTO") or "").strip(),
        "NOME_DA_PEÇA": (rec.get("NOME_DA_PEÇA") or "").strip(),
        "QTD": str(rec.get("QTD") or "").strip(),
        "CLIENTE": (rec.get("CLIENTE") or "").strip(),
        "SUPERFICIE": (rec.get("SUPERFICIE") or "").strip(),
        "LAUDO_OBSERVAÇÕES": (rec.get("LAUDO_OBSERVAÇÕES") or "").strip(),
        "DATA": normalize_data(rec.get("DATA") or ""),
        "SOLICITANTE": (rec.get("SOLICITANTE") or "").strip(),
    }
    mapa["LAUDO"] = (rec.get("LAUDO") or inferir_laudo_final(mapa["LAUDO_OBSERVAÇÕES"])).strip()
    return mapa

def gerar_docx(mapa: dict) -> str:
    modelo = MODELO_LP if mapa["TIPO"] == "LP" else MODELO_EVS
    if not os.path.exists(modelo):
        raise FileNotFoundError(f"Modelo não encontrado: {modelo}")

    doc = Document(modelo)
    aplicar_substituicoes(doc, mapa)

    os.makedirs(PASTA_OUT, exist_ok=True)

    long_name = montar_nome_longo(
        mapa["TIPO"],
        mapa["N_RELATORIO"],
        mapa["EQUIPAMENTO"],
        mapa["CLIENTE"],
        mapa.get("SOLICITANTE", ""),
        mapa["DATA"],
    )
    path_long = os.path.join(PASTA_OUT, sanitize_filename(long_name))
    doc.save(path_long)
    return path_long

def main():
    print("== GERADOR (JSON only v2.1) — nome inclui SOLICITANTE ==")
    if not os.path.exists(ARQ_JSON):
        raise FileNotFoundError(f"Não achei '{ARQ_JSON}'. Crie um input.json com a lista de relatórios.")

    with open(ARQ_JSON, "r", encoding="utf-8") as f:
        registros = json.load(f)

    print(f"Registros carregados: {len(registros)}")
    ok = falha = 0
    for i, rec in enumerate(registros, 1):
        try:
            mapa = processar_registro(rec)
            saida = gerar_docx(mapa)
            print(f"[{i}] OK -> {saida}")
            ok += 1
        except Exception as e:
            print(f"[{i}] ERRO -> {e}")
            falha += 1

    print(f"\nConcluído. Sucesso: {ok} | Falhas: {falha}")
    input("\nPressione Enter para sair...")

if __name__ == "__main__":
    main()
